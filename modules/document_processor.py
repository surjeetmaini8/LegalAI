import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

class DocumentProcessor:
    
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> str:
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n[Page {page_num + 1}]\n{page_text}"
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        
        return text
    
    def load_docx(self, file_path: str) -> str:
        
        text = ""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
        
        return text
    
    def load_txt(self, file_path: str) -> str:
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def load_document(self, file_path: str) -> Dict:
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
       
        if file_extension == '.pdf':
            text = self.load_pdf(str(file_path))
        elif file_extension == '.docx':
            text = self.load_docx(str(file_path))
        elif file_extension == '.txt':
            text = self.load_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        return {
            'source': file_path.name,
            'text': text,
            'file_path': str(file_path)
        }
    
    def load_documents_from_directory(self, directory: str) -> List[Dict]:
        
        documents = []
        supported_formats = ['.pdf', '.docx', '.txt']
        
        dir_path = Path(directory)
        if not dir_path.exists():
            print(f"Directory not found: {directory}")
            return documents
        
        for file_path in dir_path.iterdir():
            if file_path.suffix.lower() in supported_formats:
                try:
                    doc = self.load_document(str(file_path))
                    if doc['text'].strip(): 
                        documents.append(doc)
                        print(f"Loaded: {file_path.name}")
                except Exception as e:
                    print(f"Failed to load {file_path.name}: {e}")
        
        return documents
    
    def chunk_documents(self, documents: List[Dict]) -> List[LangChainDocument]:
       
        chunked_docs = []
        
        for doc in documents:
            chunks = self.text_splitter.split_text(doc['text'])
            
            for i, chunk in enumerate(chunks):
                metadata = {
                    'source': doc['source'],
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                }
                
                chunked_docs.append(
                    LangChainDocument(
                        page_content=chunk,
                        metadata=metadata
                    )
                )
        
        print(f"Created {len(chunked_docs)} chunks from {len(documents)} documents")
        return chunked_docs
    
    def process_directory(self, directory: str) -> List[LangChainDocument]:
        
        documents = self.load_documents_from_directory(directory)
        
        if not documents:
            print("No documents found to process")
            return []
        
        
        chunked_docs = self.chunk_documents(documents)
        
        return chunked_docs
